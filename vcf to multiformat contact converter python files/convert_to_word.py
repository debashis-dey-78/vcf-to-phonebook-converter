import quopri
import docx
from docx.shared import Inches, Pt

def decode_quoted_printable(text):
    try:
        return quopri.decodestring(text.encode('utf-8')).decode('utf-8')
    except:
        return text

def convert_vcf(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8', errors='ignore') as f:
        lines = f.readlines()

    contacts = []
    current_contact = {}
    
    for line in lines:
        line = line.strip()
        if line == 'BEGIN:VCARD':
            current_contact = {}
        elif line.startswith('FN:'):
            current_contact['name'] = line[3:]
        elif line.startswith('FN;'):
            if ':' in line:
                val = line.split(':', 1)[1]
                if 'QUOTED-PRINTABLE' in line:
                    val = decode_quoted_printable(val)
                current_contact['name'] = val
        elif line.startswith('TEL'):
            parts = line.split(':', 1)
            if len(parts) > 1:
                number = parts[1].strip()
                if 'numbers' not in current_contact:
                    current_contact['numbers'] = []
                if number not in current_contact['numbers']:
                    current_contact['numbers'].append(number)
        elif line == 'END:VCARD':
            if 'name' not in current_contact:
                current_contact['name'] = 'Unknown'
            if 'numbers' in current_contact:
                contacts.append(current_contact)

    contacts.sort(key=lambda x: x['name'].lower())

    doc = docx.Document()
    
    section = doc.sections[0]
    section.page_width = Inches(3)
    section.page_height = Inches(4)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1.0

    for contact in contacts:
        p = doc.add_paragraph()
        
        name_run = p.add_run(contact['name'])
        name_run.bold = True
        
        for number in contact['numbers']:
            clean_num = number.replace('-', '').replace(' ', '')
            p.add_run(f"\n{clean_num}")
            
        doc.add_paragraph()

    for _ in range(100):
        doc.add_page_break()

    doc.save(output_file)

if __name__ == '__main__':
    convert_vcf('../contacts2064.vcf', '../contacts output/phonebook.docx')
