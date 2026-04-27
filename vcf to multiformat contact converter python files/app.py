import os
import quopri
import csv
import json
import xml.etree.ElementTree as ET
import tkinter as tk
from tkinter import filedialog, messagebox, ttk

try:
    import docx
    from docx.shared import Inches, Pt
except ImportError:
    docx = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

def decode_quoted_printable(text):
    try: return quopri.decodestring(text.encode('utf-8')).decode('utf-8')
    except: return text

def parse_vcf(input_file):
    with open(input_file, 'r', encoding='utf-8', errors='ignore') as f: lines = f.readlines()
    contacts, current_contact = [], {}
    for line in lines:
        line = line.strip()
        if line == 'BEGIN:VCARD': current_contact = {}
        elif line.startswith('FN:'): current_contact['name'] = line[3:]
        elif line.startswith('FN;'):
            if ':' in line:
                val = line.split(':', 1)[1]
                if 'QUOTED-PRINTABLE' in line: val = decode_quoted_printable(val)
                current_contact['name'] = val
        elif line.startswith('TEL'):
            parts = line.split(':', 1)
            if len(parts) > 1:
                number = parts[1].strip()
                if 'numbers' not in current_contact: current_contact['numbers'] = []
                if number not in current_contact['numbers']: current_contact['numbers'].append(number)
        elif line == 'END:VCARD':
            if 'name' not in current_contact: current_contact['name'] = 'Unknown'
            if 'numbers' in current_contact: contacts.append(current_contact)
    contacts.sort(key=lambda x: x['name'].lower())
    return contacts

def rtf_escape(text):
    text = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
    res = []
    for c in text:
        if ord(c) < 128: res.append(c)
        else: res.append(f"\\u{ord(c)}?")
    return ''.join(res)

class VCFConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("VCF to Phonebook Converter")
        self.root.geometry("500x550")
        
        self.input_file = tk.StringVar()
        self.output_dir = tk.StringVar()
        
        # UI Elements
        tk.Label(root, text="Select VCF File:", font=("Arial", 10, "bold")).pack(pady=5, anchor="w", padx=20)
        file_frame = tk.Frame(root)
        file_frame.pack(fill="x", padx=20)
        tk.Entry(file_frame, textvariable=self.input_file, state="readonly", width=40).pack(side="left", padx=(0, 10))
        tk.Button(file_frame, text="Browse", command=self.browse_file).pack(side="left")

        tk.Label(root, text="Select Output Folder:", font=("Arial", 10, "bold")).pack(pady=5, anchor="w", padx=20)
        dir_frame = tk.Frame(root)
        dir_frame.pack(fill="x", padx=20)
        tk.Entry(dir_frame, textvariable=self.output_dir, state="readonly", width=40).pack(side="left", padx=(0, 10))
        tk.Button(dir_frame, text="Browse", command=self.browse_dir).pack(side="left")

        tk.Label(root, text="Select Formats to Generate:", font=("Arial", 10, "bold")).pack(pady=(15, 5), anchor="w", padx=20)
        
        self.formats = {
            "PDF (3x4 inch)": tk.BooleanVar(value=True),
            "Word (.docx) (3x4 inch)": tk.BooleanVar(value=True),
            "HTML (.html)": tk.BooleanVar(value=False),
            "Text (.txt)": tk.BooleanVar(value=False),
            "Excel (.xlsx)": tk.BooleanVar(value=False),
            "CSV (.csv)": tk.BooleanVar(value=False),
            "Markdown (.md)": tk.BooleanVar(value=False),
            "JSON (.json)": tk.BooleanVar(value=False),
            "XML (.xml)": tk.BooleanVar(value=False),
            "RTF (.rtf)": tk.BooleanVar(value=False)
        }
        
        format_frame = tk.Frame(root)
        format_frame.pack(fill="both", expand=True, padx=20)
        for fmt, var in self.formats.items():
            tk.Checkbutton(format_frame, text=fmt, variable=var).pack(anchor="w")

        self.convert_btn = tk.Button(root, text="Convert Contacts", bg="#4CAF50", fg="white", font=("Arial", 12, "bold"), command=self.convert)
        self.convert_btn.pack(pady=20)
        
        self.status_label = tk.Label(root, text="", fg="blue")
        self.status_label.pack(pady=5)

    def browse_file(self):
        filename = filedialog.askopenfilename(filetypes=[("VCF files", "*.vcf"), ("All files", "*.*")])
        if filename: self.input_file.set(filename)

    def browse_dir(self):
        dirname = filedialog.askdirectory()
        if dirname: self.output_dir.set(dirname)

    def convert(self):
        if not self.input_file.get():
            messagebox.showerror("Error", "Please select an input VCF file.")
            return
        if not self.output_dir.get():
            messagebox.showerror("Error", "Please select an output directory.")
            return
            
        selected_formats = [f for f, v in self.formats.items() if v.get()]
        if not selected_formats:
            messagebox.showerror("Error", "Please select at least one output format.")
            return

        self.status_label.config(text="Parsing VCF...")
        self.root.update()
        
        try:
            contacts = parse_vcf(self.input_file.get())
            base_path = os.path.join(self.output_dir.get(), "phonebook")
            
            for fmt in selected_formats:
                self.status_label.config(text=f"Generating {fmt}...")
                self.root.update()
                
                if fmt == "Text (.txt)": self.gen_txt(contacts, base_path + ".txt")
                elif fmt == "HTML (.html)": self.gen_html(contacts, base_path + ".html")
                elif fmt == "CSV (.csv)": self.gen_csv(contacts, base_path + ".csv")
                elif fmt == "Markdown (.md)": self.gen_md(contacts, base_path + ".md")
                elif fmt == "JSON (.json)": self.gen_json(contacts, base_path + ".json")
                elif fmt == "XML (.xml)": self.gen_xml(contacts, base_path + ".xml")
                elif fmt == "RTF (.rtf)": self.gen_rtf(contacts, base_path + ".rtf")
                elif fmt == "PDF (3x4 inch)": 
                    if FPDF: self.gen_pdf(contacts, base_path + ".pdf")
                    else: messagebox.showwarning("Missing Library", "fpdf2 is not installed. PDF generation skipped.")
                elif fmt == "Word (.docx) (3x4 inch)": 
                    if docx: self.gen_docx(contacts, base_path + ".docx")
                    else: messagebox.showwarning("Missing Library", "python-docx is not installed. Word generation skipped.")
                elif fmt == "Excel (.xlsx)": 
                    if openpyxl: self.gen_xlsx(contacts, base_path + ".xlsx")
                    else: messagebox.showwarning("Missing Library", "openpyxl is not installed. Excel generation skipped.")

            self.status_label.config(text="Conversion Complete!")
            messagebox.showinfo("Success", "All selected formats have been generated successfully.")
            
        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            self.status_label.config(text="Error occurred.")

    def gen_txt(self, contacts, out):
        with open(out, 'w', encoding='utf-8') as f:
            for c in contacts:
                f.write(f"{c['name']}\n")
                for n in c['numbers']: f.write(f"{n.replace('-', '').replace(' ', '')}\n")
                f.write("-----------------\n")

    def gen_html(self, contacts, out):
        with open(out, 'w', encoding='utf-8') as f:
            f.write("<html><head><style>@page { size: 3in 4in; margin: 0.2in; } body { font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.2; } .contact { margin-bottom: 1em; } .name { font-weight: bold; }</style></head><body>\n")
            for c in contacts:
                f.write("<div class='contact'>\n")
                f.write(f"<div class='name'>{c['name']}</div>\n")
                for n in c['numbers']: f.write(f"<div>{n.replace('-', '').replace(' ', '')}</div>\n")
                f.write("</div>\n")
            f.write("</body></html>\n")

    def gen_csv(self, contacts, out):
        with open(out, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(['Name', 'Phone 1', 'Phone 2', 'Phone 3'])
            for c in contacts: writer.writerow([c['name']] + [n.replace('-', '').replace(' ', '') for n in c['numbers']])

    def gen_md(self, contacts, out):
        with open(out, 'w', encoding='utf-8') as f:
            for c in contacts:
                f.write(f"**{c['name']}**\n")
                for n in c['numbers']: f.write(f"{n.replace('-', '').replace(' ', '')}\n")
                f.write("\n")

    def gen_json(self, contacts, out):
        with open(out, 'w', encoding='utf-8') as f:
            json.dump(contacts, f, indent=4, ensure_ascii=False)

    def gen_xml(self, contacts, out):
        root = ET.Element("Phonebook")
        for c in contacts:
            c_el = ET.SubElement(root, "Contact")
            n_el = ET.SubElement(c_el, "Name")
            n_el.text = c['name']
            nums_el = ET.SubElement(c_el, "Numbers")
            for num in c['numbers']:
                num_el = ET.SubElement(nums_el, "Number")
                num_el.text = num.replace('-', '').replace(' ', '')
        tree = ET.ElementTree(root)
        ET.indent(tree, space="  ", level=0)
        tree.write(out, encoding='utf-8', xml_declaration=True)

    def gen_rtf(self, contacts, out):
        with open(out, 'w', encoding='utf-8') as f:
            f.write("{\\rtf1\\ansi\\ansicpg1252\\deff0\\nouicompat\\deflang1033{\\fonttbl{\\f0\\fnil\\fcharset0 Arial;}}\n")
            f.write("{\\*\\generator RTF Python script;}\\viewkind4\\uc1\n")
            f.write("\\pard\\sa200\\sl276\\slmult1\\f0\\fs20\\lang9\n")
            for c in contacts:
                name = rtf_escape(c['name'])
                f.write(f"\\b {name}\\b0\\line\n")
                for number in c['numbers']: f.write(f"{number.replace('-', '').replace(' ', '')}\\line\n")
                f.write("\\line\n")
            f.write("}\n")

    def gen_pdf(self, contacts, out):
        pdf = FPDF(format=(76.2, 101.6))
        pdf.set_margins(5, 5, 5)
        pdf.set_auto_page_break(auto=True, margin=5)
        pdf.add_page()
        for contact in contacts:
            pdf.set_font("helvetica", style="B", size=10)
            name = contact['name'].encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 5, txt=name, ln=1)
            pdf.set_font("helvetica", size=10)
            for num in contact['numbers']:
                clean_num = num.replace('-', '').replace(' ', '')
                pdf.cell(0, 5, txt=clean_num, ln=1)
            pdf.ln(5)
        for _ in range(100): pdf.add_page()
        pdf.output(out)

    def gen_docx(self, contacts, out):
        doc = docx.Document()
        section = doc.sections[0]
        section.page_width = Inches(3)
        section.page_height = Inches(4)
        section.left_margin, section.right_margin = Inches(0.2), Inches(0.2)
        section.top_margin, section.bottom_margin = Inches(0.2), Inches(0.2)
        style = doc.styles['Normal']
        style.font.name, style.font.size = 'Arial', Pt(10)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        for contact in contacts:
            p = doc.add_paragraph()
            name_run = p.add_run(contact['name'])
            name_run.bold = True
            for number in contact['numbers']:
                clean_num = number.replace('-', '').replace(' ', '')
                p.add_run(f"\n{clean_num}")
            doc.add_paragraph()
        for _ in range(100): doc.add_page_break()
        doc.save(out)

if __name__ == "__main__":
    root = tk.Tk()
    app = VCFConverterApp(root)
    root.mainloop()
